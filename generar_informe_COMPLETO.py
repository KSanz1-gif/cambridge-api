import json
from datetime import date
from docx import Document

# Cargar plantilla
with open("plantilla_informe.json", "r", encoding="utf-8") as f:
    base = json.load(f)["estructura"]

# Datos simulados para rellenar
base["nombre_estudiante"] = "Juan Pérez"
base["nivel_simulado"] = "B2"
base["fecha"] = str(date.today())
base["puntuacion_estimativa"] = {
    "Grammatical Range and Accuracy": "Good control of tenses, minor slips",
    "Lexical Resource": "Sufficient range for topic, some repetition",
    "Discourse Management": "Coherent speech with some hesitation",
    "Pronunciation": "Mostly intelligible, slight L1 influence",
    "Interactive Communication": "Responds appropriately, minor delays"
}
base["errores_detectados"] = {
    "grammar": ["Confused past simple and present perfect"],
    "vocabulary": ["Used 'do a mistake' instead of 'make a mistake'"],
    "pronunciation": ["Mispronounced 'think' as 'sink'"]
}
base["comentario_general"] = "A strong B2 performance with clear effort and control. Minor issues in grammar and pronunciation."
base["recomendaciones"] = [
    "Review use of present perfect.",
    "Practise common collocations.",
    "Work on /θ/ and /ð/ pronunciation."
]

# Crear documento Word
doc = Document()
doc.add_heading(f"Informe de Speaking - {base['nombre_estudiante']}", 0)
doc.add_paragraph(f"Nivel simulado: {base['nivel_simulado']}")
doc.add_paragraph(f"Fecha: {base['fecha']}")
doc.add_heading("Puntuación estimativa", level=1)

for criterio, valor in base["puntuacion_estimativa"].items():
    doc.add_paragraph(f"{criterio}: {valor}", style="List Bullet")

doc.add_heading("Errores detectados", level=1)
for categoria, errores in base["errores_detectados"].items():
    doc.add_paragraph(f"{categoria.capitalize()}:")
    for error in errores:
        doc.add_paragraph(f"- {error}", style="List Bullet")

doc.add_heading("Comentario general", level=1)
doc.add_paragraph(base["comentario_general"])

doc.add_heading("Recomendaciones", level=1)
for rec in base["recomendaciones"]:
    doc.add_paragraph(f"- {rec}", style="List Bullet")

# Guardar documento
doc.save("informe_speaking_juan.docx")
print("Informe generado correctamente.")